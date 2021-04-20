# Author: Kirill Zhosul.
# Description: Python program for working with diplomas in these extensions - *.docx. This
# program takes data in these extensions - *.xlsx format and converts it to output file extension.

# Importing modules for correct working.
import openpyxl  # Module for working with documents (*.xlsx).
import os  # Module for working with paths to files.
import re  # Module for working with regular expressions.
import docx  # Module for working with documents (.docx)
from PIL import Image  # Module for working with images.
from PIL import ImageFont  # Module for working with fonts.
from PIL import ImageDraw  # Module for working with drawing.
from time import time  # Module for checking completion time.


class Client:
    # Class client for working with client row as class instance.

    def __init__(self, cell_values):
        # Class constructor.

        # List with all cell values.
        self.__cell_values = cell_values

        # I think, this shouldn't looks like that by IDK RN.
        self.type = self.__cell_values[0]
        self.name = self.__cell_values[1]
        self.school = self.__cell_values[2]
        self.givenfor = self.__cell_values[3]
        self.event = self.__cell_values[4]
        self.date = self.__cell_values[5]

    def __str__(self):
        # Class str() override method.

        # Return cell values list as string.
        return str(self.__cell_values)


def document_create_png(document_path, document_client):
    """
    Creates png document.
    :param document_path: Document path for creating.
    :param document_client: Client object for getting data.
    :return:
    """
    if not enable_overwrite:
        # If overwrite is not enabled.
        if os.path.exists(document_path + image_extension):
            # If file already exists.

            # Returning void
            return

    # Opening sample as image.
    try:
        sample_image = Image.open(sample_filename)
    except FileNotFoundError:
        # Showing an error.
        print("Ошибка! Не удалось найти файл шаблона *.png.")
        # Exiting program.
        raise SystemExit

    # Getting image W, H.
    if 'image_w' in globals():
        image_w, image_h = sample_image.size

    # Opening image as draw.
    sample_draw = ImageDraw.Draw(sample_image)

    # Getting font.

    # Drawing document type (Diploma / Certificate)
    text = document_client.type
    text_font = ImageFont.truetype(font="arialbd.ttf", size=72)
    text_w, text_h = sample_draw.textsize(text, font=text_font)
    sample_draw.text(((image_w - text_w) / 2, (image_h - text_h) / 2 - image_h / 6), text, (79, 127, 222),
                     font=text_font)

    # Drawing awarded text
    text = "награждается"
    text_font = ImageFont.truetype(font="arialbd.ttf", size=16)
    text_w, text_h = sample_draw.textsize(text, font=text_font)
    sample_draw.text(((image_w - text_w) / 2, (image_h - text_h) / 2 - image_h / 8), text, (0, 0, 0), font=text_font)

    # Saving image.
    sample_image.save(document_path + ".png")


def document_create_docx(document_path):
    """
    Creates docx document.
    :param document_path: Document path for creating.
    :return:
    """
    if not enable_overwrite:
        # If overwrite is not enabled.
        if os.path.exists(document_path + document_extension):
            # If file already exists.

            # Returning void
            return

    # Creating new document.
    document = docx.Document()

    # Adding image.
    document.add_paragraph().add_run().add_picture(document_path + image_extension, docx.shared.Inches(1.25))

    # Saving document.
    document.save(document_path + document_extension)


def create_documents(document_path, document_client):
    """
    Creates all packet of documents.
    :param document_path: Document path for creating.
    :param document_client: Client object for getting data.
    :return:
    """

    # Creating documents.
    document_create_png(document_path, document_client)
    document_create_docx(document_path)


def xlsx_get_cell(current_sheet, cell_row, cell_number):
    """
    Function for getting cell value from xlsx sheet.
    :param current_sheet: Current sheet from which get cell.
    :param cell_row: Cell row like A (A1)
    :param cell_number: Cell number like 1 (A1)
    :return: String or none
    """

    # Returning cell.
    return current_sheet[f"{cell_row}{cell_number}"].value


def xlsx_check_row(current_sheet, row_number):
    """
    Function that checks row for having values in all cells.
    :param current_sheet: Current sheet from which get cell.
    :param row_number: Row number like 1 (A1)
    :return:
    """
    for row_symbol in ["A", "B", "C", "D", "E", "F"]:
        # Iterating over all symbols.
        if xlsx_get_cell(current_sheet, row_symbol, row_number) is None:
            # If we don`t have value in this cell.

            # Returning false.
            return False
    # True if we don`t returned False above.
    return True


if __name__ == "__main__":
    # Entry point of the program.

    # Time when we start execution of our program.
    time_start = time()

    image_w = None
    image_h = None

    # WIP Temporary! Input filenames with excel table.
    input_extension = ".xlsx"
    image_extension = ".png"
    input_filename = os.getcwd() + '\\documents\\in\\clients' + input_extension
    sample_filename = os.getcwd() + '\\documents\\in\\template' + image_extension
    out_directory = os.getcwd() + '\\documents\\out\\'
    out_directory_rule = "{name}"
    out_filename_rule = "{name} {type} {givenfor} ({date})"
    document_extension = ".docx"
    # Should code overwrite old files or just create new one.
    enable_overwrite = True

    # Opening workbook from xlsx table.
    if input_filename.endswith(".xlsx"):
        # If it is XLSX format.

        # Opening workbook.
        workbook = openpyxl.load_workbook(input_filename)

        # Getting worksheet (Last active).
        # TODO: Add selecting of sheet.
        worksheet = workbook.active

        # Current index for getting values.
        current_index = 2  # 2 As first row is just for human editor, just marking what in row, skip it.

        while True:
            # Endless loop, breaking when if-statement is called.

            if xlsx_check_row(worksheet, current_index):
                # If keys with this index is existing in worksheet.

                # All cells values in row as an list.
                row_values = []

                for row_symbol in ["A", "B", "C", "D", "E", "F"]:
                    # Iterating over all symbols.
                    row_values.append(xlsx_get_cell(worksheet, row_symbol, current_index))

                if len(re.findall(r'[\w]+', row_values[5], re.U)) == 0:
                    # If this row contain only trash symbols.

                    # Go to next line and not parse this.
                    current_index += 1
                    continue

                # Creating new client instance for working with it.
                client = Client(row_values)

                # Saving image as out image file.
                try:
                    # Formatting directory.
                    current_filename = out_directory_rule.format(type=client.type, name=client.name,
                                                                 school=client.school, givenfor=client.givenfor,
                                                                 event=client.event, date=client.date) + "\\"

                    # Making directory if not exists.
                    if not os.path.exists(out_directory + current_filename):
                        os.mkdir(out_directory + current_filename)

                    # Formatting filename.
                    current_filename += out_filename_rule.format(type=client.type, name=client.name,
                                                                 school=client.school, givenfor=client.givenfor,
                                                                 event=client.event, date=client.date)
                except KeyError:
                    # Showing an error.
                    print("Ошибка! Неккоректное правило имени выходного файла.")
                    # Exiting program.
                    raise SystemExit

                # Creating documents.
                create_documents(out_directory + current_filename, client)

                # Increasing current index for iterating over all current indices.
                current_index += 1

                # Continue to next cycle.
                continue
            # Exiting endless loop if if-statement above is not completed.
            break

        # Showing result time for function.
        print(f"Обработано {current_index - 2 - 1} человек за  {round(time() - time_start, 3)}c")
    else:
        # Showing an error.
        print("Ошибка! Неизвестный формат входного файла.")
        # Exiting program.
        raise SystemExit
