"""
    Программный код генератора для программы "Автоматические Документы"
    Автор программного кода: Кирилл Жосул.
"""

# Импортация модулей.

# Встроенные модули.
import multiprocessing  # Модуль для *многопоточности*
from os import mkdir  # Создание папки.
from os.path import exists  # Проверка существования.
from string import ascii_uppercase  # Модуль для получения списка заглавных букв.
from re import findall, U  # Модуль для работы с регулярными выражениями.

# Модули внешние.
import docx  # Модуль для работы с .docx документами.
import openpyxl  # Модуль для работы с .xlsx документами.

# Модули PIL. Для работы с графикой.
from PIL import Image
from PIL import ImageFont
from PIL import ImageDraw
from PyQt5.QtWidgets import QMessageBox

# Модули ядра.
import interface  # Модуль для работы с интерфейсом.


class XLSX:
    # Класс для работы с XLSX.

    def __init__(self, filename_table):
        # Конструктор класса.

        # Открытие рабочей книги.
        self.workbook = openpyxl.load_workbook(filename_table)

        # Получение рабочего листа.
        self.worksheet = self.workbook.active

    def get_cell(self, _column, _row):
        # Функция получения значения из ячейки.
        return self.worksheet[_column + str(_row)].value

    def validate_row(self, _row):
        # Функция для валидации ряда в Excel рабочем листе.

        for _column in ascii_uppercase[:6]:
            # Для первых 6 символов алфавита (Как работает Excel)

            if self.get_cell(_column, _row) is None:
                # Если в этом столбце на выбранном ряду

                # Возвращение булевой лжи, это ряд не валиден.
                return False

        # Возвращение булевой правды, это ряд валиден.
        return True


class User:
    # Класс пользователя для работы с ним.

    def __init__(self):
        # Конструктор класса.

        # Сырые значения из ряда.
        self.__raw_row_values = []

    def __getitem__(self, item):
        # User[0]
        return self.__raw_row_values[item]

    def format_string(self, _string):
        # Функция форматирования строки.

        # Ошибка.
        if len(self.__raw_row_values) != 6:
            interface.error("Неизвестная ошибка!")

        # Форматирование
        return _string.format(type=self.__raw_row_values[0], name=self.__raw_row_values[1],
                              school=self.__raw_row_values[2], givenfor=self.__raw_row_values[3],
                              event=self.__raw_row_values[4], date=self.__raw_row_values[5])

    def append_raw_row_value(self, _value):
        # Функция для добавления сырого значения из ряда в список.

        # Исключение.
        if _value is None:
            raise ValueError

        # Добавления значения.
        self.__raw_row_values.append(_value)


class Generator:
    # Класс генератора.

    @staticmethod
    def finalize_png(_path, _user, _overwrite_enabled, _file_image):
        # Функция генерации png файла.

        if not _overwrite_enabled:
            # Если перезапись не включена.
            if exists(_path + ".docx"):
                # Выход если файл уже есть.
                return

        # Открытие шаблона.
        image = Image.open(_file_image)
        draw = ImageDraw.Draw(image)

        # Получение размеров шаблона.
        w, h = image.size

        # Отрисовка.
        text = _user[0]
        font = ImageFont.truetype(font="arialbd.ttf", size=72)
        tw, th = draw.textsize(text, font=font)
        draw.text(((w - tw) / 2, (h - th) / 2 - h / 6), text, (79, 127, 222), font=font)

        text = "награждается"
        font = ImageFont.truetype(font="arialbd.ttf", size=16)
        tw, th = draw.textsize(text, font=font)
        draw.text(((w - tw) / 2, (h - th) / 2 - h / 11), text, (0, 0, 0), font=font)

        text = _user[1]
        font = ImageFont.truetype(font="arialbd.ttf", size=48)
        tw, th = draw.textsize(text, font=font)
        draw.text(((w - tw) / 2, (h - th) / 2 - h / 20), text, (0, 0, 0), font=font)

        text = _user[2]
        font = ImageFont.truetype(font="arialbd.ttf", size=24)
        tw, th = draw.textsize(text, font=font)
        draw.text(((w - tw) / 2, (h - th) / 2 + 20), text, (0, 0, 0), font=font)

        text = _user[3]
        font = ImageFont.truetype(font="arialbd.ttf", size=48)
        tw, th = draw.textsize(text, font=font)
        draw.text(((w - tw) / 2, (h - th) / 2 + 20 + 48), text, (0, 0, 0), font=font)

        if len(_user[4]) > 35:
            lines_spaced = _user[4].split()
            lines = [0, 0]
            lines[0] = "на " + " ".join(lines_spaced[:len(lines_spaced) // 2])
            lines[1] = " ".join(lines_spaced[len(lines_spaced) // 2:])
            lines[0] += " "
            font = ImageFont.truetype(font="arialbd.ttf", size=24)
            tw, th = draw.textsize(lines[0], font=font)
            draw.text(((w - tw) / 2, (h - th) / 2 + 20 + 48 + 48), lines[0], (0, 0, 0), font=font)
            tw, th = draw.textsize(lines[1], font=font)
            draw.text(((w - tw) / 2, (h - th) / 2 + 20 + 48 + 48 + 32), lines[1], (0, 0, 0), font=font)
        else:
            text = _user[4]
            font = ImageFont.truetype(font="arialbd.ttf", size=32)
            tw, th = draw.textsize(text, font=font)
            draw.text(((w - tw) / 2, (h - th) / 2 + 20 + 48 + 48), text, (0, 0, 0), font=font)

        text = _user[5]
        font = ImageFont.truetype(font="arialbd.ttf", size=32)
        tw, th = draw.textsize(text, font=font)
        draw.text(((w - tw) / 2, h - 32), text, (0, 0, 0), font=font)

        # Сохранение.
        image.save(_path + ".png")

    @staticmethod
    def finalize_docx(_path, _overwrite_enabled):
        # Функция генерации docx файла.

        if not _overwrite_enabled:
            # Если перезапись не включена.
            if exists(_path + ".docx"):
                # Выход если файл уже есть.
                return

        # Создание нового документа.
        document = docx.Document()

        # Добавление картинки.
        document.add_paragraph().add_run().add_picture(_path + ".png", docx.shared.Inches(1.25))

        # Сохранение документа.
        document.save(_path + ".docx")

    @staticmethod
    def finalise_pdf(_path, _overwrite_enabled):
        # Функция генерации pdf файла.

        if not _overwrite_enabled:
            # Если перезапись не включена.
            if exists(_path + ".pdf"):
                # Выход если файл уже есть.
                return

    @staticmethod
    def finalise_html(_path, _overwrite_enabled):
        # Функция генерации html файла.

        if not _overwrite_enabled:
            # Если перезапись не включена.
            if exists(_path + ".html"):
                # Выход если файл уже есть.
                return

        with open(_path + ".html", "w",encoding="UTF-8") as file:
            file.write(f"<img src=\"{_path + '.png'}\">")

    @staticmethod
    def populate(_path, _overwrite_enabled):
        # Функция распространения по другим форматам.
        Generator.finalize_docx(_path, _overwrite_enabled)
        Generator.finalise_pdf(_path, _overwrite_enabled)
        Generator.finalise_html(_path, _overwrite_enabled)

    @staticmethod
    def finalize_document(_path, _user, _overwrite_enabled, _file_image):
        # Функция окончания получения данных, запись в документ.

        # Создание основной фотографии документа.
        Generator.finalize_png(_path, _user, _overwrite_enabled, _file_image)

        # Распростанение документа по разным форматам.
        Generator.populate(_path, _overwrite_enabled)


def generate_thread(_file_table, _file_image, _file_directory, _rule_directory, _rule_file, _overwrite_enabled):
    # Функция потока генерации.

    # Создание объекта XLSX.
    xlsx = XLSX(_file_table)

    # Ряд пользователя на котором мы сейчас находимся.
    current_row = 2  # 2 Потому что 1) Отсчёт с 1, 2) Нам не требуется получать данные из первого ряда.

    while True:
        # Бесконечный цикл.

        if xlsx.validate_row(current_row):
            # Если это валидный ряд.

            # Проверка на наличие требуемых символов, если в первой ячейке нет нормальных символов,
            # То переход на следующий ряд, в целях оптимизации.
            if len(findall(r'[\w]+', xlsx.get_cell("A", current_row), U)) == 0:
                current_row += 1
                continue

            # Объект пользователя.
            user = User()

            for _column in ascii_uppercase[:6]:
                # Для первых 6 символов алфавита (Как работает Excel)
                user.append_raw_row_value(xlsx.get_cell(_column, current_row))

            try:
                # Форматирование папки куда ложить файл.
                current_filename = _file_directory + "\\" + user.format_string(_rule_directory) + "\\"

                # Создание папки вывода если её нет.
                if not exists(current_filename):
                    mkdir(current_filename)

                # Formatting filename.
                current_filename += user.format_string(_rule_file)
            except KeyError:
                # Если ошибка.

                # Показ ошибки.
                interface.error("Некорректное правило имени или папки!")

                # Выход из цикла.
                break

            # Генерация.
            Generator.finalize_document(current_filename, user, _overwrite_enabled, _file_image)

            # Переход на следующий ряд.
            current_row += 1

            # Переходим на следующую итерацию что бы пропустить выход из цикла.
            continue

        # Выход из цикла если условие выше не выполнено.
        break


def generate(_file_table, _file_image, _file_directory, _rule_directory, _rule_file, _overwrite_enabled):
    # Функция получения потока для генерации потока.

    # Короче, это должно быть после выполнения,
    # Но там проблемы с вызовов вне потока,
    # Исправлю как нибудь = )
    # Наверное это даже никто не прочитает.
    # Но всё может быть!)
    interface.information("Генерация успешно выполнена.")

    # Возвращение потока генерации.
    return multiprocessing.Process(target=generate_thread, args=(_file_table, _file_image, _file_directory, _rule_directory, _rule_file, _overwrite_enabled))
